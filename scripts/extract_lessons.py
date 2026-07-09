#!/usr/bin/env python3
"""Converts the "Daily Neonatology Lessons" .docx series into per-lesson JSON
for the NeoRef Learn tab. Reads from a source directory (not part of this
repo), writes one JSON file per lesson into public/lessons/, plus an index
(lessons-index.json) with day/chapter/title/authors metadata for all lessons.

Usage: python scripts/extract_lessons.py <source_dir> <content_out_dir> <index_out_path>
"""
import json
import re
import sys
from pathlib import Path

import docx
from docx.oxml.ns import qn

FILENAME_RE = re.compile(
    r"^Day\s+(\d+)\s*[-–]?\s*(Avery|Fanaroff)\s+Ch(\d+)\s+(.+)$"
)


def parse_filename(stem):
    m = FILENAME_RE.match(stem)
    if not m:
        return None
    day, book, chapter, title = m.groups()
    return {
        "day": int(day),
        "book": book,
        "chapter": int(chapter),
        "titleFromFilename": title.strip(),
    }


def cell_text(cell):
    return "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip())


def extract_body(doc):
    """Walks the document body in true order, splitting into a title block
    (everything before the first Heading 1) and a list of content blocks."""
    body = doc.element.body
    paragraphs_by_id = {id(p._p): p for p in doc.paragraphs}
    tables_by_id = {id(t._tbl): t for t in doc.tables}

    title_lines = []
    blocks = []
    seen_first_heading = False

    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            p = paragraphs_by_id.get(id(child))
            if p is None:
                continue
            text = p.text.strip()
            style = p.style.name if p.style else "Normal"
            if not text:
                continue
            if style == "Heading 1":
                seen_first_heading = True
                blocks.append({"type": "h1", "text": text})
            elif style == "Heading 2":
                seen_first_heading = True
                blocks.append({"type": "h2", "text": text})
            elif not seen_first_heading:
                title_lines.append(text)
            elif style == "List Paragraph":
                blocks.append({"type": "li", "text": text})
            else:
                blocks.append({"type": "p", "text": text})
        elif child.tag == qn("w:tbl"):
            t = tables_by_id.get(id(child))
            if t is None:
                continue
            rows = [[cell_text(c) for c in row.cells] for row in t.rows]
            rows = [r for r in rows if any(c for c in r)]
            if not rows:
                continue
            if len(rows) == 1 and len(rows[0]) == 1:
                blocks.append({"type": "callout", "text": rows[0][0]})
            else:
                blocks.append({"type": "table", "rows": rows})

    return title_lines, blocks


def main():
    src_dir = Path(sys.argv[1])
    out_dir = Path(sys.argv[2])
    index_out_path = Path(sys.argv[3])
    out_dir.mkdir(parents=True, exist_ok=True)

    index = []
    skipped = []

    for path in sorted(src_dir.glob("*.docx")):
        if path.name.startswith("~$"):
            continue
        meta = parse_filename(path.stem)
        if meta is None:
            skipped.append(path.name)
            continue

        doc = docx.Document(str(path))
        title_lines, blocks = extract_body(doc)

        # title_lines pattern (consistent across the series): [0]="DAILY
        # NEONATOLOGY — DAY N", [1]="<Book>'s ... (Year)", [2]=chapter title,
        # [3]=byline, sometimes suffixed with "— Avery 11th ed. 2024".
        authors = ""
        if len(title_lines) > 3:
            authors = title_lines[3].split("—")[0].strip()

        lesson = {
            "day": meta["day"],
            "book": meta["book"],
            "chapter": meta["chapter"],
            "title": meta["titleFromFilename"],
            "authors": authors,
            "blocks": blocks,
        }

        out_path = out_dir / f"day-{meta['day']:03d}.json"
        out_path.write_text(json.dumps(lesson, ensure_ascii=False), encoding="utf-8")

        index.append(
            {
                "day": meta["day"],
                "book": meta["book"],
                "chapter": meta["chapter"],
                "title": meta["titleFromFilename"],
                "authors": authors,
            }
        )

    index.sort(key=lambda l: l["day"])
    index_out_path.write_text(
        json.dumps(index, ensure_ascii=False, indent=1), encoding="utf-8"
    )

    print(f"Converted {len(index)} lessons, skipped {len(skipped)}: {skipped}")


if __name__ == "__main__":
    main()
